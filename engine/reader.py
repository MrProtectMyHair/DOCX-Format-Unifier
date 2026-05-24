"""DOCX 读取模块 — 提取段落和表格的结构化信息"""

from docx import Document


def read_docx(filepath):
    """读取 DOCX 文件，返回 DocData 结构。

    返回格式:
    {
        "paragraphs": [ParaInfo, ...],
        "tables": [TableInfo, ...],
    }
    """
    doc = Document(filepath)
    return {
        "paragraphs": _extract_paragraphs(doc),
        "tables": _extract_tables(doc),
    }


def _extract_paragraphs(doc):
    """提取所有段落的文本和格式信息。"""
    result = []
    for para in doc.paragraphs:
        text = para.text
        if text is None:
            text = ""
        fmt = para.paragraph_format
        alignment = fmt.alignment
        line_spacing = fmt.line_spacing
        space_before = fmt.space_before
        space_after = fmt.space_after

        # 获取第一个非空 run 的字符格式
        run = _first_meaningful_run(para)
        font_name = run.font.name if (run and run.font.name) else None
        font_size = run.font.size if (run and run.font.size) else None
        bold = run.font.bold if run else None
        italic = run.font.italic if run else None

        result.append({
            "text": text,
            "alignment": alignment,
            "font_name": font_name,
            "font_size": font_size,
            "bold": bold,
            "italic": italic,
            "line_spacing": line_spacing,
            "space_before": space_before,
            "space_after": space_after,
        })
    return result


def _first_meaningful_run(para):
    """返回段落中第一个有字体信息的 run，用于获取段落格式。"""
    for run in para.runs:
        if run.text and run.text.strip():
            return run
    if para.runs:
        return para.runs[0]
    return None


def _extract_tables(doc):
    """提取所有表格的结构和单元格内容。"""
    result = []
    for table in doc.tables:
        rows = len(table.rows)
        cols = _effective_column_count(table)

        cells = []
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                text = cell.text if cell.text else ""
                cells.append({
                    "row": ri,
                    "col": ci,
                    "text": text,
                })

        result.append({
            "rows": rows,
            "cols": cols,
            "cells": cells,
        })
    return result


def _effective_column_count(table):
    """计算表格的有效列数（考虑合并单元格）。

    通过检查每行的实际 cell 数量取最大值。
    python-docx 对合并单元格的 GridSpan 支持不完善，
    这里使用行内最大 cell 数作为列数。
    """
    max_cols = 0
    for row in table.rows:
        cell_count = len(row.cells)
        if cell_count > max_cols:
            max_cols = cell_count
    return max_cols
