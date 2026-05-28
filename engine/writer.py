"""DOCX 生成模块 — 以模板为骨架，替换文字内容"""

import os
import shutil
import tempfile
import uuid
from docx import Document


def generate_output(template_path, input_data, template_data, para_matches, table_matches, unmatched_paras, output_path):
    """生成输出文件。

    策略：复制模板到临时文件（避免 Windows Explorer 锁），修改后移到目标路径。
    """
    # 1. 复制模板到临时文件（用 uuid 生成路径，不预先创建文件，避免杀软锁占位文件）
    tmp_path = os.path.join(tempfile.gettempdir(), f'tfu_{uuid.uuid4().hex}.docx')
    try:
        shutil.copy2(template_path, tmp_path)
    except (IOError, PermissionError) as e:
        raise IOError("无法写入输出文件，请检查文件是否被其他程序占用或路径是否有写入权限") from e

    # 2. 打开副本进行修改
    doc = Document(tmp_path)

    # 3. 替换段落文本
    _replace_paragraphs(doc, input_data, template_data, para_matches, unmatched_paras)

    # 4. 替换表格单元格文本（遍历所有表格）
    for ti, tm in enumerate(table_matches):
        if ti < len(doc.tables) and ti < len(input_data["tables"]) and ti < len(template_data["tables"]):
            _replace_table_cells(doc.tables[ti], input_data["tables"][ti],
                                template_data["tables"][ti], tm["cell_matches"])

    # 4.5. 空值填充（遍历所有表格）
    for ti, tm in enumerate(table_matches):
        if ti < len(doc.tables) and ti < len(input_data["tables"]) and ti < len(template_data["tables"]):
            _fill_empty_cells_by_label(doc.tables[ti], input_data["tables"][ti],
                                       template_data["tables"][ti], tm["cell_matches"])

    # 5. 保存到临时文件
    try:
        doc.save(tmp_path)
    except (IOError, PermissionError) as e:
        os.unlink(tmp_path)
        raise IOError("无法保存输出文件，请关闭已打开的输出文件后重试") from e

    # 6. 移动到目标路径
    try:
        shutil.move(tmp_path, output_path)
    except (IOError, PermissionError) as e:
        os.unlink(tmp_path)
        raise IOError("无法写入输出文件，请检查文件是否被其他程序占用或路径是否有写入权限") from e


def _replace_paragraphs(doc, input_data, template_data, para_matches, unmatched_paras):
    """替换模板段落中的文本为输入文件的文本。

    当一个输入段落匹配到多个模板段落时（复合段落），按标签位置分割文本。
    """
    input_paras = input_data["paragraphs"]
    tmpl_paras = template_data["paragraphs"]
    from engine.matcher import _extract_label

    # 建立: 模板索引 → 输入索引
    tmpl_to_input = {ti: ii for ii, ti in para_matches}

    # 检测复合段落：同一个输入索引对应多个模板段落
    input_to_tmpls = {}
    for ii, ti in para_matches:
        if ii not in input_to_tmpls:
            input_to_tmpls[ii] = []
        input_to_tmpls[ii].append(ti)

    # 计算切分映射: (input_idx, template_idx) -> text_subset
    _split_map = {}
    for ii, ti_list in input_to_tmpls.items():
        if len(ti_list) <= 1:
            continue
        text = input_paras[ii]["text"]
        label_positions = []
        for ti2 in ti_list:
            lbl = _extract_label(tmpl_paras[ti2]["text"])
            if lbl:
                pos = text.find(lbl)
                if pos >= 0:
                    label_positions.append((pos, lbl, ti2))
        label_positions.sort()
        for idx, (pos, lbl, ti2) in enumerate(label_positions):
            start = pos
            if idx + 1 < len(label_positions):
                end = label_positions[idx + 1][0]
            else:
                end = len(text)
            _split_map[(ii, ti2)] = text[start:end].strip()

    # 未匹配段落的索引（后续会从中移除被合并的）
    unmatched_indices = [i for i in unmatched_paras if input_paras[i]["text"].strip()]

    for ti, para in enumerate(doc.paragraphs):
        if ti in tmpl_to_input:
            ii = tmpl_to_input[ti]
            if (ii, ti) in _split_map:
                new_text = _split_map[(ii, ti)]
            else:
                new_text = input_paras[ii]["text"]
            # 检查未匹配段落：若其 label 在模板中存在但不在当前 input 中，合并进来
            for u_i in list(unmatched_indices):
                u_label = _extract_label(input_paras[u_i]["text"])
                if u_label and len(u_label) >= 2:
                    if u_label in tmpl_paras[ti]["text"] and u_label not in new_text:
                        new_text += input_paras[u_i]["text"]
                        unmatched_indices.remove(u_i)
            # 如果输入段落含「标签：值」格式但模板段落只有值（无标签），剥离标签
            new_text = _strip_label_if_template_has_none(
                new_text, tmpl_paras[ti]["text"])
            _fill_form_paragraph(para, new_text)

    # 追加未匹配的输入段落
    for ii in unmatched_indices:
        new_text = input_paras[ii]["text"]
        last_para = doc.paragraphs[-1] if doc.paragraphs else None
        new_para = doc.add_paragraph(new_text)
        in_para = input_paras[ii]
        if in_para["font_name"]:
            for run in new_para.runs:
                run.font.name = in_para["font_name"]
        if in_para["font_size"]:
            for run in new_para.runs:
                run.font.size = in_para["font_size"]


def _replace_table_cells(table, in_tbl, tmpl_tbl, cell_matches):
    """替换模板表格单元格中的文本为输入文件的文本。"""
    if not cell_matches:
        return

    # 建立输入单元格的文本查找: (row, col) → text
    in_cell_map = {}
    for cell in in_tbl["cells"]:
        in_cell_map[(cell["row"], cell["col"])] = cell["text"]

    # 给模板每个单元格分配顺序 ID（与 reader 遍历顺序相同）
    tmpl_pos_to_id = {}
    tmpl_id = 0
    for cell in tmpl_tbl["cells"]:
        key = (cell["row"], cell["col"])
        if key not in tmpl_pos_to_id:
            tmpl_pos_to_id[key] = tmpl_id
            tmpl_id += 1

    # 建立: 模板顺序 ID → 替换文本
    id_mapping = {}
    for m in cell_matches:
        in_key = (m["input_row"], m["input_col"])
        tmpl_key = (m["tmpl_row"], m["tmpl_col"])
        if in_key in in_cell_map and tmpl_key in tmpl_pos_to_id:
            cell_id = tmpl_pos_to_id[tmpl_key]
            id_mapping[cell_id] = in_cell_map[in_key]

    # 按相同遍历顺序在输出文档中替换
    out_id = 0
    for row in table.rows:
        for cell in row.cells:
            if out_id in id_mapping:
                # 跳过签章/行政类单元格
                if not _has_skip_keyword(cell.text):
                    _set_cell_text(cell, id_mapping[out_id])
            out_id += 1


def _set_paragraph_text(para, text):
    """设置段落的文本，保留模板的字符格式。"""
    # 保留第一个 run 的格式，用其承载全部文本
    if para.runs:
        # 将所有文本合并到第一个 run，清空其余 run
        first_run = para.runs[0]
        first_run.text = text
        for run in para.runs[1:]:
            run.text = ""
    else:
        # 没有 run 时添加一个
        run = para.add_run(text)


def _is_blank_run(run):
    """判断 run 是否为表单空白占位 — 纯下划线或带下划线格式的纯空白。"""
    if not run.text:
        return False
    if all(c == '_' for c in run.text):
        return True
    if not run.text.strip() and run.font.underline:
        return True
    return False


def _fill_form_paragraph(para, input_text):
    """表单式段落填充：保留空白 run 的格式，从输入文本提取值填入。

    对模板段落中的每个 blank run（下划线/空白标记），用固定文本作为锚点
    从输入文本中定位并提取对应值，保留原 run 的字体和下划线格式。
    """
    runs = list(para.runs)
    if not runs:
        _set_paragraph_text(para, input_text)
        return

    # 1. 构建骨架 — 分类 run 并合并连续空白
    skeleton = []  # {'type':'fixed'/'slot', 'runs':[idx,...] for slot, 'run':idx for fixed, 'text':str}
    i = 0
    while i < len(runs):
        if _is_blank_run(runs[i]):
            slot_runs = [i]
            i += 1
            while i < len(runs) and _is_blank_run(runs[i]):
                slot_runs.append(i)
                i += 1
            skeleton.append({'type': 'slot', 'runs': slot_runs})
        else:
            skeleton.append({'type': 'fixed', 'run': i, 'text': runs[i].text})
            i += 1

    has_slot = any(s['type'] == 'slot' for s in skeleton)
    has_fixed = any(s['type'] == 'fixed' and s['text'].strip() for s in skeleton)
    if not has_slot or not has_fixed:
        # 无空白 run：比较去空格文本，一致则保留模板 run 结构
        import re
        tmpl_text = ''.join(r.text for r in runs)
        if re.sub(r'\s+', '', tmpl_text) == re.sub(r'\s+', '', input_text):
            return
        _set_paragraph_text(para, input_text)
        return

    # 2. 在输入文本中搜索每个固定文本 → 锚点位置 (run_idx, start, end)
    fixed_entries = [(s['run'], s['text'].strip()) for s in skeleton
                     if s['type'] == 'fixed' and s['text'].strip()]
    anchors = []
    scan_pos = 0
    import re
    for run_idx, ft in fixed_entries:
        pos = input_text.find(ft, scan_pos)
        match_end = pos + len(ft) if pos >= 0 else None
        if pos < 0 and len(re.sub(r'\s+', '', ft)) >= 2:
            # 空格不一致回退：用无空格版做模糊匹配
            ft_ns = re.sub(r'\s+', '', ft)
            pattern = re.escape(ft_ns[0]) + ''.join(
                r'\s*' + re.escape(c) for c in ft_ns[1:])
            m = re.search(pattern, input_text[scan_pos:])
            if m:
                pos = scan_pos + m.start()
                match_end = scan_pos + m.end()
        if pos >= 0:
            anchors.append((run_idx, pos, match_end))
            scan_pos = match_end

    # 3. 计算每个槽的填充值
    slot_entries = [(s['runs'], s) for s in skeleton if s['type'] == 'slot']
    fill_values = {}  # first_run_idx → text

    for slot_indices, _ in slot_entries:
        first_run = slot_indices[0]

        prev = None
        for a in anchors:
            if a[0] < first_run:
                prev = a
        nxt = None
        for a in anchors:
            if a[0] > first_run:
                nxt = a
                break

        start = prev[2] if prev else 0
        end = nxt[1] if nxt else len(input_text)
        value = input_text[start:end]

        # 最后一个槽：若后方还有固定文本未匹配到，用其首字符截断
        if nxt is None and slot_indices == slot_entries[-1][0]:
            trailing = ''
            for s in skeleton:
                if s['type'] == 'fixed' and s['run'] > first_run:
                    trailing += s['text']
            trailing = trailing.lstrip()
            if trailing:
                for pl in range(1, len(trailing) + 1):
                    p = value.find(trailing[:pl])
                    if p >= 0:
                        value = value[:p]
                        break

        fill_values[first_run] = value

    # 4. 应用：空白 run 用填充值替换，保留格式；多余空白 run 置空
    for s in skeleton:
        if s['type'] == 'fixed':
            pass  # 保留模板固定 run
        else:
            slot_runs = s['runs']
            value = fill_values.get(slot_runs[0], '')
            runs[slot_runs[0]].text = value
            for r in slot_runs[1:]:
                runs[r].text = ''


def _set_cell_text(cell, text):
    """设置单元格的文本，保留模板的字符格式。

    按换行符拆分为多段，匹配模板的段落结构。
    多余的空段落从 XML 中物理移除，避免空行占高。
    """
    # 按换行拆分，过滤空字符串（连续换行不产生空段落）
    raw_parts = text.split("\n") if text else [""]
    parts = [p for p in raw_parts if p.strip()]
    if not parts:
        parts = [""]
    n_paras = len(cell.paragraphs)

    # 填充前 N 个段落（N = min(len(parts), n_paras)）
    for i, part in enumerate(parts):
        if i < n_paras:
            _set_paragraph_text(cell.paragraphs[i], part)
        else:
            # 输入文本的段数超过模板段落数，在末尾段落追加
            _set_paragraph_text(cell.paragraphs[-1],
                               cell.paragraphs[-1].text + "\n" + part)

    # 移除多余的空段落（物理删除，而非仅置空）
    for i in range(len(parts), n_paras):
        para = cell.paragraphs[i]
        p_elem = para._element
        p_elem.getparent().remove(p_elem)


# 签章/行政类标签 — 其右侧空单元格不需要填值
_SKIP_LABELS = {"签字", "盖章", "审批", "审核", "日期", "负责人", "经办人", "分管"}


def _has_skip_keyword(text):
    """判断文本是否包含签章/行政类关键词。"""
    for kw in _SKIP_LABELS:
        if kw in text:
            return True
    return False


def _is_summary_label(text):
    """判断文本是否包含汇总类关键词（合计/总计/小计）。"""
    for kw in ("合计", "总计", "小计"):
        if kw in text:
            return True
    return False


def _fill_empty_cells_by_label(table, in_tbl, tmpl_tbl, cell_matches):
    """标签领地填充：每个匹配标签管辖其到下一个标签之间的空单元格。

    领地内所有空单元格填入该标签在输入侧的第一个值 (d_col=1)。
    遇到输入值越界或撞到其他标签时，回退到最近的有效值。
    """
    if not cell_matches:
        return

    # 输入单元格查找
    in_cell_map = {}
    for cell in in_tbl["cells"]:
        in_cell_map[(cell["row"], cell["col"])] = cell["text"]

    # 模板单元格查找
    tmpl_cell_map = {}
    for cell in tmpl_tbl["cells"]:
        tmpl_cell_map[(cell["row"], cell["col"])] = cell["text"]

    # 模板顺序 ID 映射
    tmpl_pos_to_id = {}
    seq = 0
    for cell in tmpl_tbl["cells"]:
        k = (cell["row"], cell["col"])
        if k not in tmpl_pos_to_id:
            tmpl_pos_to_id[k] = seq
            seq += 1

    # 已匹配标签: template (r,c) → input (r,c)
    label_map = {}
    for m in cell_matches:
        label_map[(m["tmpl_row"], m["tmpl_col"])] = (m["input_row"], m["input_col"])

    # 收集每行的已匹配标签列
    row_labels = {}  # row -> sorted list of (col, input_label_pos)
    for (tr, tc), in_pos in label_map.items():
        row_labels.setdefault(tr, []).append((tc, in_pos))
    for tr in row_labels:
        row_labels[tr].sort()

    # 按标签领地填充（同行向左找标签）
    for tr, labels in row_labels.items():
        for idx, (tc, in_pos) in enumerate(labels):
            tmpl_text = tmpl_cell_map.get((tr, tc), "")
            if _has_skip_keyword(tmpl_text):
                continue

            if idx + 1 < len(labels):
                territory_end = labels[idx + 1][0]
            else:
                territory_end = max(c for (r, c) in tmpl_cell_map if r == tr) + 1

            # 汇总类标签（合计/总计/小计）仅填紧邻一格，不扩展领地
            if _is_summary_label(tmpl_text):
                territory_end = min(territory_end, tc + 2)

            in_val_key = (in_pos[0], in_pos[1] + 1)
            if in_val_key not in in_cell_map:
                continue
            in_val = in_cell_map[in_val_key].strip()
            if not in_val or in_val_key in label_map.values():
                for fallback_d in range(2, 5):
                    fb_key = (in_pos[0], in_pos[1] + fallback_d)
                    if fb_key in in_cell_map and fb_key not in label_map.values():
                        in_val = in_cell_map[fb_key].strip()
                        if in_val:
                            break
            if not in_val:
                continue

            for fill_c in range(tc + 1, territory_end):
                _fill_one_cell(table, tmpl_cell_map, tmpl_pos_to_id, tr, fill_c, in_val, tc)

    # 补充：同行没有标签的空单元格，向上查找表头
    _fill_by_column_header(table, tmpl_cell_map, tmpl_pos_to_id, label_map, in_cell_map)


def _fill_one_cell(table, tmpl_cell_map, tmpl_pos_to_id, tr, tc, in_val, label_col=None):
    """填入单个单元格的值，可选从标签列复制格式。"""
    fill_key = (tr, tc)
    if fill_key not in tmpl_cell_map:
        return
    if tmpl_cell_map[fill_key].strip():
        return
    if _has_skip_keyword(tmpl_cell_map.get((tr, tc), "")):
        return
    cell_id = tmpl_pos_to_id.get(fill_key)
    if cell_id is None:
        return
    out_id = 0
    for row in table.rows:
        for cell in row.cells:
            if out_id == cell_id:
                _set_cell_text(cell, in_val)
                if label_col is not None:
                    _copy_label_format(row, label_col, tc)
            out_id += 1


def _fill_by_column_header(table, tmpl_cell_map, tmpl_pos_to_id, label_map, in_cell_map):
    """列头引导填充：仅处理「全空数据行」，不干扰领地填充。

    先找到上方最近的「表头行」（该行有>=2个已匹配标签），
    然后按顺序偏移从输入查找对应数据行，按列映射填入。
    """
    if not label_map:
        return

    max_row = max(r for (r, c) in tmpl_cell_map) if tmpl_cell_map else 0

    # 找到所有有标签的行
    rows_with_labels = set(tr for (tr, tc) in label_map)

    # 收集需要保护的模板行（含签字等关键词）
    skip_rows = set()
    for (tr, tc), text in tmpl_cell_map.items():
        if _has_skip_keyword(text):
            skip_rows.add(tr)

    # 处理全空数据行（该行没有已匹配标签）
    for tr in range(max_row + 1):
        if tr in rows_with_labels:
            continue
        if tr in skip_rows:
            continue

        # 找到该空行上方最近的「表头行」（至少有2个标签的行）
        header_row = None
        for scan_r in range(tr - 1, -1, -1):
            n = sum(1 for (sr, sc) in label_map if sr == scan_r)
            if n >= 2:
                header_row = scan_r
                break
        if header_row is None:
            continue

        # 从该表头行构建列映射
        col_map = {}
        in_header_row = None
        for (sr, sc), (ir, ic) in label_map.items():
            if sr == header_row:
                if in_header_row is None:
                    in_header_row = ir
                if sc not in col_map:
                    col_map[sc] = ic
        if in_header_row is None:
            continue

        # 回退：未匹配的表头列，用低阈值尝试匹配输入列
        if col_map:
            from engine.matcher import _extract_label, _label_similarity
            all_tmpl_cols = set(tc for (r, tc) in tmpl_cell_map if r == header_row)
            all_in_cols = set(c for (r, c) in in_cell_map if r == in_header_row)
            for tc in all_tmpl_cols:
                if tc in col_map:
                    continue
                tmpl_text = tmpl_cell_map.get((header_row, tc), "")
                if not tmpl_text.strip():
                    continue
                tmpl_lbl = _extract_label(tmpl_text)
                best_ic = None
                best_score = 0.0
                for ic in all_in_cols:
                    in_text = in_cell_map.get((in_header_row, ic), "")
                    in_lbl = _extract_label(in_text)
                    score = _label_similarity(tmpl_lbl, in_lbl)
                    if score > best_score and score >= 0.5:
                        best_score = score
                        best_ic = ic
                if best_ic is not None:
                    col_map[tc] = best_ic
                    all_in_cols.discard(best_ic)

        if not col_map:
            continue

        # 按顺序偏移找输入数据行
        offset = tr - header_row
        in_data_row = in_header_row + offset

        for tc in col_map:
            fill_key = (tr, tc)
            if fill_key not in tmpl_cell_map:
                continue
            if tmpl_cell_map[fill_key].strip():
                continue

            in_col = col_map[tc]
            in_val_key = (in_data_row, in_col)
            if in_val_key not in in_cell_map:
                continue
            in_val = in_cell_map[in_val_key].strip()
            if not in_val or in_val_key in label_map.values():
                continue

            _fill_one_cell(table, tmpl_cell_map, tmpl_pos_to_id, tr, tc, in_val, tc)


def _copy_label_format(row, label_col, target_col):
    """将同行标签单元格的字体格式复制到目标单元格。"""
    try:
        label_cell = row.cells[label_col] if label_col < len(row.cells) else None
        target_cell = row.cells[target_col] if target_col < len(row.cells) else None
        if not label_cell or not target_cell:
            return
        if not label_cell.paragraphs or not target_cell.paragraphs:
            return
        src_run = None
        for r in label_cell.paragraphs[0].runs:
            if r.text and r.text.strip():
                src_run = r
                break
        if not src_run and label_cell.paragraphs[0].runs:
            src_run = label_cell.paragraphs[0].runs[0]
        if not src_run:
            return
        dst_run = None
        for r in target_cell.paragraphs[0].runs:
            if r.text and r.text.strip():
                dst_run = r
                break
        if not dst_run and target_cell.paragraphs[0].runs:
            dst_run = target_cell.paragraphs[0].runs[0]
        if not dst_run:
            dst_run = target_cell.paragraphs[0].add_run("")
        # 复制格式
        if src_run.font.name:
            dst_run.font.name = src_run.font.name
        if src_run.font.size:
            dst_run.font.size = src_run.font.size
        if src_run.font.bold is not None:
            dst_run.font.bold = src_run.font.bold
        # 段落对齐
        src_align = label_cell.paragraphs[0].paragraph_format.alignment
        if src_align is not None:
            target_cell.paragraphs[0].paragraph_format.alignment = src_align
    except Exception:
        pass  # 格式复制失败不阻塞主流程


def _strip_label_if_template_has_none(input_text, template_text):
    """如果输入含「标签：值」但模板不含该标签，则剥离输入的标签部分。

    例：输入 "填表及制发日期: 2026.4.1"，模板 "2026.4.1"
    → 返回 "2026.4.1"
    """
    from engine.matcher import _extract_label
    in_label = _extract_label(input_text)

    if not in_label or len(in_label) < 2:
        return input_text
    # 检查模板是否包含输入的标签（作为独立词出现）
    # 同时检查去空格版本，因为 _extract_label 会去掉所有空格
    import re
    tmpl_no_space = re.sub(r'\s+', '', template_text)
    if in_label not in template_text and in_label not in tmpl_no_space:
        # 再检查模板 label 的核心部分（去下划线）是否出现在输入 label 中
        # 例：模板"______专家级别：____" → core="专家级别" ∈ 输入"阮骏琳专家级别"
        tmpl_label = _extract_label(template_text)
        if tmpl_label:
            tmpl_core = re.sub(r'[\s_]+', '', tmpl_label)
            if len(tmpl_core) >= 2 and tmpl_core not in in_label:
                m = re.search(r'[:：]\s*(.+)', input_text)
                if m:
                    return m.group(1).strip()
        else:
            m = re.search(r'[:：]\s*(.+)', input_text)
            if m:
                return m.group(1).strip()
    return input_text
